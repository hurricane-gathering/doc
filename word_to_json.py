import json
import re
import docx
from typing import Dict, List, Any, Optional


def is_chinese(text: str) -> bool:
    """Check if the text contains Chinese characters."""
    return bool(re.search(r'[\u4e00-\u9fff]', text))


def clean_text(text: str) -> str:
    """Remove leading bullets and whitespace."""
    # Match bullets like •, -, *, etc.
    text = re.sub(r'^[\s•\-\*]+', '', text).strip()
    return text


def parse_date_range(text: str):
    """Parse 'Start - End' string."""
    parts = text.split('-')
    if len(parts) >= 2:
        return parts[0].strip(), parts[-1].strip()
    return text.strip(), ""


def get_empty_multilingual() -> Dict[str, str]:
    return {"zh": "", "en": ""}


def set_multilingual(target: Dict[str, str], text: str):
    if is_chinese(text):
        target["zh"] = text
    else:
        # Or assume en if not zh, or put in primary language bucket.
        target["en"] = text
        # Given requirement: "adapt one language, other returns empty string"
        # If text is Chinese, put in zh. If text is English, put in en.
        # But if the whole doc is Chinese, maybe even English words should go to 'zh' fields?
        # Let's assume language is per-document.
        pass


class ResumeParser:
    def __init__(self, file_path: str):
        self.doc = docx.Document(file_path)
        self.data = {
            "fontSize": {"zh": 10.5, "en": 11},  # Default/Placeholder
            "lineHeight": {"zh": 1.1, "en": 1},  # Default/Placeholder
            "personalInfo": {
                "avatar": None,
                "firstName": get_empty_multilingual(),
                "lastName": get_empty_multilingual(),
                "email": "",
                "phone": "",
                "linkedin": ""
            },
            "education": [],
            "experience": [],
            "activities": [],
            "skillsInterests": {
                "languages": [],
                "skills": get_empty_multilingual(),
                "interests": []
            }
        }
        self.current_lang = "zh"  # Default to zh, will detect

    def parse(self) -> Dict[str, Any]:
        # Simple extraction of font/line height from first paragraph if available
        if self.doc.paragraphs:
            p = self.doc.paragraphs[0]
            if p.runs and p.runs[0].font.size:
                # pt = p.runs[0].font.size.pt
                # self.data["fontSize"][self.current_lang] = pt
                pass

        # Determine main language from first few lines
        full_text = " ".join([p.text for p in self.doc.paragraphs[:5]])
        self.current_lang = "zh" if is_chinese(full_text) else "en"

        paragraphs = [p.text.strip()
                      for p in self.doc.paragraphs if p.text.strip()]

        # Indexes for sections
        sections = {
            "education": ["教育背景", "Education"],
            "experience": ["实习经历", "Work Experience", "Experience"],
            "activities": ["课外活动经历", "Activities", "Leadership"],
            "skills": ["技能/兴趣", "Skills & Interests", "Skills", "Interests"]
        }

        section_indices = {}
        for i, text in enumerate(paragraphs):
            for key, keywords in sections.items():
                if any(k in text for k in keywords):
                    section_indices[key] = i
                    break

        # Sort indices to know where sections end
        sorted_indices = sorted(section_indices.items(), key=lambda x: x[1])

        # Parse Personal Info (Before first section)
        first_section_idx = sorted_indices[0][1] if sorted_indices else len(
            paragraphs)
        self._parse_personal_info(paragraphs[:first_section_idx])

        # Parse Sections
        for i, (section_name, start_idx) in enumerate(sorted_indices):
            end_idx = sorted_indices[i+1][1] if i + \
                1 < len(sorted_indices) else len(paragraphs)
            content = paragraphs[start_idx+1: end_idx]

            if section_name == "education":
                self._parse_education(content)
            elif section_name == "experience":
                self._parse_experience(content)
            elif section_name == "activities":
                self._parse_activities(content)
            elif section_name == "skills":
                self._parse_skills(content)

        return self.data

    def _set_field(self, obj, key, value):
        """Helper to set zh/en field based on current language."""
        if self.current_lang == "zh":
            obj[key]["zh"] = value
        else:
            obj[key]["en"] = value

    def _parse_personal_info(self, lines):
        if not lines:
            return

        # Name (First line)
        name = lines[0]
        # Basic splitting for name
        if self.current_lang == 'zh':
            # Assume Surname Firstname for Chinese usually, or just put whole in Last Name or split?
            # User example: firstName: "姓", lastName: "名".
            # If "张三", firstName="张", lastName="三"
            if len(name) >= 2:
                self._set_field(
                    self.data["personalInfo"], "firstName", name[0])
                self._set_field(
                    self.data["personalInfo"], "lastName", name[1:])
            else:
                self._set_field(self.data["personalInfo"], "lastName", name)
        else:
            parts = name.split()
            if parts:
                self._set_field(
                    self.data["personalInfo"], "firstName", parts[0])
                self._set_field(self.data["personalInfo"],
                                "lastName", " ".join(parts[1:]))

        # Contact (Second line)
        if len(lines) > 1:
            contact_line = lines[1]
            # Extract phone - allowing 'x' for placeholders as seen in example
            phone_match = re.search(
                r'[\+\d\-\(\)\sxxx]{8,}', contact_line, re.IGNORECASE)
            if phone_match:
                self.data["personalInfo"]["phone"] = phone_match.group().strip()

            # Extract email
            email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', contact_line)
            if email_match:
                self.data["personalInfo"]["email"] = email_match.group().strip()

            # LinkedIn - check remaining text or "LinkedIn" keyword
            if "linkedin" in contact_line.lower():
                # Placeholder or extract URL
                self.data["personalInfo"]["linkedin"] = "LinkedIn"

    def _parse_education(self, lines):
        entry = None

        for line in lines:
            # Check if line is Institution line (has tab or looks like a header)
            # Heuristic: Tab separates School and Date
            if '\t' in line:
                if entry:
                    self.data["education"].append(entry)

                parts = line.split('\t')
                school = parts[0].strip()
                date_str = parts[1].strip() if len(parts) > 1 else ""
                start, end = parse_date_range(date_str)

                entry = {
                    "institution": get_empty_multilingual(),
                    "major": get_empty_multilingual(),
                    "startDate": start,
                    "endDate": end,
                    "gpa": "",
                    "gpaRanking": "",
                    "relevantCourses": get_empty_multilingual(),
                    "honors": get_empty_multilingual()
                }
                self._set_field(entry, "institution", school)
                continue

            if entry is None:
                continue

            # Degree line (usually follows institution, no bullet)
            if not line.strip().startswith('•') and not entry["major"][self.current_lang]:
                self._set_field(entry, "major", line.strip())
                continue

            # Bullets
            clean_line = clean_text(line)
            if "GPA" in line:
                # GPA:Your Score/4.0 (前xx%)
                gpa_match = re.search(r'GPA[:\s]*([^\(]+)', clean_line)
                if gpa_match:
                    entry["gpa"] = gpa_match.group(1).strip()

                rank_match = re.search(r'\(([^)]+)\)', clean_line)
                if rank_match:
                    entry["gpaRanking"] = rank_match.group(1).strip()

            elif "课程" in line or "Courses" in line:
                # 核心课程：...
                content = re.sub(r'.*[:：]', '', clean_line).strip()
                self._set_field(entry, "relevantCourses", content)

            elif "荣誉" in line or "Honors" in line:
                content = re.sub(r'.*[:：]', '', clean_line).strip()
                self._set_field(entry, "honors", content)

        if entry:
            self.data["education"].append(entry)

    def _parse_experience(self, lines):
        # Similar to education but with company, position, bullets
        entry = None
        for line in lines:
            if '\t' in line and not line.strip().startswith('•'):
                # Heuristic: if it has tab and isn't a bullet, it's likely a header line.
                # Experience often has 2 header lines: Company+Date, then Position+Location
                parts = line.split('\t')

                # Check if this is Company line or Position line
                # Company line usually has Date range
                has_date = re.search(r'\d{4}|Year|Month', parts[-1])

                if has_date:
                    # Save previous
                    if entry:
                        self.data["experience"].append(entry)

                    company = parts[0].strip()
                    date_str = parts[1].strip() if len(parts) > 1 else ""
                    start, end = parse_date_range(date_str)

                    entry = {
                        "company": get_empty_multilingual(),
                        "position": get_empty_multilingual(),
                        "location": get_empty_multilingual(),
                        "startDate": start,
                        # Schema varies: education endDate is string, experience is object? User schema shows experience endDate is object.
                        "endDate": {"zh": end, "en": ""} if self.current_lang == 'zh' else {"zh": "", "en": end},
                        "actionPoints": [],
                        "isOpen": "Present" in end or "至今" in end or not end
                    }
                    self._set_field(entry, "company", company)
                else:
                    # Likely Position + Location
                    if entry:
                        position = parts[0].strip()
                        location = parts[1].strip() if len(parts) > 1 else ""
                        self._set_field(entry, "position", position)
                        self._set_field(entry, "location", location)
                continue

            if entry is None:
                continue

            # Bullets (Action Points)
            if line.strip().startswith('•') or line.strip().startswith('-'):
                summary = clean_text(line)
                point = {
                    "taskAction": get_empty_multilingual(),
                    "approach": get_empty_multilingual(),
                    "implementationProcess": get_empty_multilingual(),
                    "quantifiableResults": get_empty_multilingual(),
                    "impact": get_empty_multilingual(),
                    "summary": get_empty_multilingual()
                }
                self._set_field(point, "summary", summary)
                entry["actionPoints"].append(point)

        if entry:
            self.data["experience"].append(entry)

    def _parse_activities(self, lines):
        # Similar structure to experience in user query
        entry = None
        for line in lines:
            if '\t' in line:
                parts = line.split('\t')
                # Company/Org line
                if re.search(r'\d{4}|Year|Month', parts[-1]):
                    if entry:
                        self.data["activities"].append(entry)

                    org = parts[0].strip()
                    date_str = parts[1].strip() if len(parts) > 1 else ""
                    start, end = parse_date_range(date_str)

                    entry = {
                        "company": get_empty_multilingual(),  # Schema uses "company" for activities too
                        "position": get_empty_multilingual(),
                        "location": get_empty_multilingual(),
                        "startDate": start,
                        "endDate": {"zh": end, "en": ""} if self.current_lang == 'zh' else {"zh": "", "en": end},
                        "actionPoints": [],
                        "isOpen": False
                    }
                    self._set_field(entry, "company", org)
                else:
                    # Role line?
                    # The example shows "角色1" on a new line without tab sometimes?
                    # Let's check explore_docx output: P30: '角色1'. No tab.
                    # P29 has tab.
                    # So if tab not found, it might be role line if entry exists and position is empty
                    pass
            elif entry and not line.strip().startswith('•') and not entry["position"][self.current_lang]:
                self._set_field(entry, "position", line.strip())

            if entry:
                if line.strip().startswith('•'):
                    summary = clean_text(line)
                    point = {
                        "taskAction": get_empty_multilingual(),
                        "approach": get_empty_multilingual(),
                        "implementationProcess": get_empty_multilingual(),
                        "quantifiableResults": get_empty_multilingual(),
                        "impact": get_empty_multilingual(),
                        "summary": get_empty_multilingual()
                    }
                    self._set_field(point, "summary", summary)
                    entry["actionPoints"].append(point)

        if entry:
            self.data["activities"].append(entry)

    def _parse_skills(self, lines):
        # Languages, Skills, Interests
        for line in lines:
            clean = clean_text(line)
            if "语言" in line or "Languages" in line:
                content = re.sub(r'.*[:：]', '', clean).strip()
                # Split by comma or semicolon
                langs = re.split(r'[，,、]', content)
                for lang in langs:
                    if not lang.strip():
                        continue
                    # Extract proficiency if in parens e.g. "英语（母语）"
                    prof = "native"  # Default
                    lang_name = lang.strip()
                    match = re.match(r'([^\(（]+)[\(（]([^\)）]+)[\)）]', lang)
                    if match:
                        lang_name = match.group(1).strip()
                        prof_text = match.group(2).strip()
                        # Map prof text to schema keys if needed, or just keep text?
                        # Schema has "proficiency": "native".
                        pass

                    l_obj = {"language": get_empty_multilingual(),
                             "proficiency": prof}
                    self._set_field(l_obj, "language", lang_name)
                    self.data["skillsInterests"]["languages"].append(l_obj)

            elif "技能" in line or "Skills" in line:
                content = re.sub(r'.*[:：]', '', clean).strip()
                self._set_field(
                    self.data["skillsInterests"], "skills", content)

            elif "兴趣" in line or "Interests" in line:
                content = re.sub(r'.*[:：]', '', clean).strip()
                # Split by comma/semicolon
                # Example: 兴趣爱好1（达成的成就1）、兴趣爱好2（达成的成就2）
                interests = re.split(r'[，,、]', content)
                for item in interests:
                    if not item.strip():
                        continue
                    # Extract achievement
                    match = re.match(r'([^\(（]+)[\(（]([^\)）]+)[\)）]', item)
                    name = item.strip()
                    achievement = ""
                    if match:
                        name = match.group(1).strip()
                        achievement = match.group(2).strip()

                    i_obj = {
                        "interest": get_empty_multilingual(),
                        "achievement": get_empty_multilingual()
                    }
                    self._set_field(i_obj, "interest", name)
                    self._set_field(i_obj, "achievement", achievement)
                    self.data["skillsInterests"]["interests"].append(i_obj)


def parse_docx_to_json(file_path):
    parser = ResumeParser(file_path)
    return parser.parse()


if __name__ == "__main__":
    result = parse_docx_to_json("test_output.docx")
    print(json.dumps(result, indent=2, ensure_ascii=False))

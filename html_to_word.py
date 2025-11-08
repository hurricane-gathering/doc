from pathlib import Path
from typing import Optional, Union, Tuple
import re
import copy
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def html_to_word(
    html_content: str,
    output_path: Optional[Union[str, Path]] = None
) -> Document:
    """将 HTML 转换为 Word 文档

    Args:
        html_content: HTML 字符串
        output_path: 输出路径，为 None 则不保存

    Returns:
        Document 对象

    Raises:
        ValueError: HTML 内容为空
        IOError: 文件保存失败
    """

    doc = Document()
    soup = BeautifulSoup(html_content, 'lxml')

    # 从 HTML 中提取字号和行距
    font_size_pt, line_height = _extract_css_styles(soup)

    # 设置文档样式（使用提取的样式）
    _setup_styles(doc, font_size_pt, line_height)

    # 添加姓名和联系信息
    _process_header(doc, soup, font_size_pt)

    # 添加间隔
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 处理所有章节
    _process_sections(doc, soup, font_size_pt, line_height)

    # 保存文档
    if output_path:
        _save_document(doc, output_path)

    return doc


def _extract_css_styles(soup: BeautifulSoup) -> Tuple[float, float]:
    """从 HTML 中提取 CSS 样式（字号和行距）

    Args:
        soup: BeautifulSoup 对象

    Returns:
        (font_size_pt, line_height) 元组
        font_size_pt: 字号（单位：磅），默认 11
        line_height: 行距倍数，默认 1.15
    """
    font_size_pt = 11.0  # 默认值
    line_height = 1.15  # 默认值

    # 查找所有 style 标签
    style_tags = soup.find_all('style')
    for style_tag in style_tags:
        css_content = style_tag.string or ''

        # 查找 .a4-page 类的样式
        # 匹配模式：.a4-page { ... font-size: 10.5pt; ... line-height: 1.6; ... }
        a4_page_match = re.search(
            r'\.a4-page\s*\{[^}]*\}',
            css_content,
            re.DOTALL | re.IGNORECASE
        )

        if a4_page_match:
            css_block = a4_page_match.group(0)

            # 提取 font-size
            font_size_match = re.search(
                r'font-size\s*:\s*([\d.]+)\s*(pt|px|em|rem)?',
                css_block,
                re.IGNORECASE
            )
            if font_size_match:
                size_value = float(font_size_match.group(1))
                size_unit = (font_size_match.group(2) or 'pt').lower()

                # 转换为磅（pt）
                if size_unit == 'pt':
                    font_size_pt = size_value
                elif size_unit == 'px':
                    font_size_pt = size_value * 0.75  # 1px ≈ 0.75pt
                elif size_unit == 'em' or size_unit == 'rem':
                    font_size_pt = size_value * 12  # 假设基础字号为 12pt
                else:
                    font_size_pt = size_value

            # 提取 line-height
            line_height_match = re.search(
                r'line-height\s*:\s*([\d.]+)',
                css_block,
                re.IGNORECASE
            )
            if line_height_match:
                line_height = float(line_height_match.group(1))

    return font_size_pt, line_height


def _setup_styles(doc: Document, font_size_pt: float, line_height: float):
    """配置文档默认样式

    Args:
        doc: Document 对象
        font_size_pt: 字号（磅）
        line_height: 行距倍数
    """
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(font_size_pt)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    # 设置默认行距
    style.paragraph_format.line_spacing = line_height


def _process_header(doc: Document, soup: BeautifulSoup, base_font_size: float):
    """处理姓名和联系信息

    Args:
        doc: Document 对象
        soup: BeautifulSoup 对象
        base_font_size: 基础字号（磅）
    """
    # 添加姓名
    name_header = soup.find('div', class_='name-header')
    if name_header:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(name_header.get_text(strip=True))
        # 姓名字号使用基础字号的 2.38 倍（约 25pt，当基础为 10.5pt 时）
        run.font.size = Pt(base_font_size * 2.38)
        run.bold = True
        para.paragraph_format.space_after = Pt(6)

    # 添加联系信息
    contact_div = soup.find('div', class_='contact-info')
    if contact_div:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(contact_div.get_text(strip=True))
        # 联系信息字号使用基础字号的 1.14 倍（约 12pt，当基础为 10.5pt 时）
        run.font.size = Pt(base_font_size * 1.14)
        para.paragraph_format.space_after = Pt(12)


def _process_sections(doc: Document, soup: BeautifulSoup, font_size_pt: float, line_height: float):
    """处理所有章节

    Args:
        doc: Document 对象
        soup: BeautifulSoup 对象
        font_size_pt: 字号（磅）
        line_height: 行距倍数
    """
    sections = soup.find_all('div', class_='section-title')

    for idx, section in enumerate(sections):
        # 添加章节标题
        _add_section_title(doc, section.get_text(strip=True), font_size_pt)

        # 处理该章节的列表
        next_ul = section.find_next_sibling('ul', class_='ul-section')
        if next_ul:
            _process_list(doc, next_ul, font_size_pt, line_height)

        # 章节间距（最后一个除外）
        if idx < len(sections) - 1:
            doc.add_paragraph().paragraph_format.space_after = Pt(8)


def _add_section_title(doc: Document, title: str, font_size_pt: float):
    """添加章节标题（加粗 + 底部边框）

    Args:
        doc: Document 对象
        title: 标题文本
        font_size_pt: 字号（磅）
    """
    para = doc.add_paragraph()
    run = para.add_run(title)
    # 章节标题字号使用基础字号的 1.24 倍（约 13pt，当基础为 10.5pt 时）
    run.font.size = Pt(font_size_pt * 1.24)
    run.bold = True

    # 添加底部边框
    p = para._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # 设置间距
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.space_before = Pt(6)


def _process_list(doc: Document, ul_element, font_size_pt: float, line_height: float):
    """处理列表元素

    Args:
        doc: Document 对象
        ul_element: ul 元素
        font_size_pt: 字号（磅）
        line_height: 行距倍数
    """
    for li in ul_element.find_all('li', recursive=False):
        has_dot = li.find('span', class_='dot') is not None

        # 提取右对齐文本
        right_span = li.find('span', class_='right-span')
        right_text = right_span.get_text(strip=True) if right_span else ""
        if right_span:
            right_span.extract()

        # 移除圆点元素（但保留其位置信息）
        dot_span = li.find('span', class_='dot')
        if dot_span:
            dot_span.extract()

        # 添加列表项，支持部分加粗
        _add_list_item_with_formatting(
            doc, li, has_dot, right_text, font_size_pt, line_height)


def _add_list_item(doc: Document, text: str, is_bold: bool, is_italic: bool,
                   has_dot: bool, right_text: str):
    """添加单个列表项（保留用于兼容性）"""
    para = doc.add_paragraph()

    # 添加圆点（如果需要）
    if has_dot:
        para.paragraph_format.left_indent = Inches(0.3)
        para.add_run('• ').font.size = Pt(11)

    # 添加主文本
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.bold = is_bold
    run.italic = is_italic

    # 添加右对齐文本（如果有）
    if right_text:
        right_run = para.add_run('\t' + right_text)
        right_run.font.size = Pt(11)
        right_run.bold = is_bold
        right_run.italic = is_italic
        para.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.0), WD_ALIGN_PARAGRAPH.RIGHT)

    # 设置紧凑的行距
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.line_spacing = 1.15


def _add_list_item_with_formatting(doc: Document, li_element, has_dot: bool, right_text: str,
                                   font_size_pt: float, line_height: float):
    """添加单个列表项，支持部分加粗和斜体

    Args:
        doc: Document 对象
        li_element: li 元素
        has_dot: 是否有圆点
        right_text: 右对齐文本
        font_size_pt: 字号（磅）
        line_height: 行距倍数
    """
    para = doc.add_paragraph()

    # 添加圆点（如果需要）
    if has_dot:
        para.paragraph_format.left_indent = Inches(0.3)
        para.add_run('• ').font.size = Pt(font_size_pt)

    # 遍历 li 元素的所有子节点，分别处理加粗和非加粗部分
    for element in li_element.children:
        if isinstance(element, Tag):
            # 处理标签元素
            if element.name == 'b' or element.name == 'strong':
                # 加粗文本，移除换行符但保留空格
                text = element.get_text(separator=' ', strip=False)
                text = re.sub(r'[\n\r]+', ' ', text)  # 将换行符替换为空格
                text = text.strip()
                if text:
                    run = para.add_run(text)
                    run.font.size = Pt(font_size_pt)
                    run.bold = True
            elif element.name == 'i' or element.name == 'em':
                # 斜体文本，移除换行符但保留空格
                text = element.get_text(separator=' ', strip=False)
                text = re.sub(r'[\n\r]+', ' ', text)
                text = text.strip()
                if text:
                    run = para.add_run(text)
                    run.font.size = Pt(font_size_pt)
                    run.italic = True
            elif element.name in ['span', 'div', 'p']:
                # 递归处理嵌套元素
                _process_nested_elements(para, element, font_size_pt)
            else:
                # 其他标签，提取文本，移除换行符但保留空格
                text = element.get_text(separator=' ', strip=False)
                text = re.sub(r'[\n\r]+', ' ', text)
                text = text.strip()
                if text:
                    run = para.add_run(text)
                    run.font.size = Pt(font_size_pt)
        elif isinstance(element, NavigableString):
            # 处理纯文本节点，移除换行符但保留空格
            text = str(element)
            # 将换行符替换为空格，然后移除首尾空白
            text = re.sub(r'[\n\r]+', ' ', text)
            text = text.strip()
            if text:
                run = para.add_run(text)
                run.font.size = Pt(font_size_pt)

    # 添加右对齐文本（如果有）
    if right_text:
        right_run = para.add_run('\t' + right_text)
        right_run.font.size = Pt(font_size_pt)
        para.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.0), WD_ALIGN_PARAGRAPH.RIGHT)

    # 设置行距（使用从 CSS 提取的值）
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.line_spacing = line_height


def _process_nested_elements(para, element, font_size_pt: float):
    """递归处理嵌套元素，保持格式

    Args:
        para: 段落对象
        element: 元素对象
        font_size_pt: 字号（磅）
    """
    for child in element.children:
        if isinstance(child, Tag):
            if child.name == 'b' or child.name == 'strong':
                # 加粗文本，移除换行符但保留空格
                text = child.get_text(separator=' ', strip=False)
                text = re.sub(r'[\n\r]+', ' ', text)
                text = text.strip()
                if text:
                    run = para.add_run(text)
                    run.font.size = Pt(font_size_pt)
                    run.bold = True
            elif child.name == 'i' or child.name == 'em':
                # 斜体文本，移除换行符但保留空格
                text = child.get_text(separator=' ', strip=False)
                text = re.sub(r'[\n\r]+', ' ', text)
                text = text.strip()
                if text:
                    run = para.add_run(text)
                    run.font.size = Pt(font_size_pt)
                    run.italic = True
            elif child.name in ['span', 'div', 'p']:
                _process_nested_elements(para, child, font_size_pt)
            else:
                # 其他标签，移除换行符但保留空格
                text = child.get_text(separator=' ', strip=False)
                text = re.sub(r'[\n\r]+', ' ', text)
                text = text.strip()
                if text:
                    run = para.add_run(text)
                    run.font.size = Pt(font_size_pt)
        elif isinstance(child, NavigableString):
            # 处理纯文本节点，移除换行符但保留空格
            text = str(child)
            # 将换行符替换为空格，然后移除首尾空白
            text = re.sub(r'[\n\r]+', ' ', text)
            text = text.strip()
            if text:
                run = para.add_run(text)
                run.font.size = Pt(font_size_pt)


def _save_document(doc: Document, output_path: Union[str, Path]):
    """保存文档"""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        doc.save(str(output_path))
        print(f"✓ Word 文档已保存至: {output_path}")
    except Exception as e:
        raise IOError(f"保存文档失败: {e}")
